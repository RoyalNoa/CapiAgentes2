"""
Ruta: ia_workspace/agentes/capi_desktop/handler.py
Descripción: Handler principal del Capi Desktop
Estado: Activo
Autor: Claude Code
Última actualización: 2025-01-14
Referencias: AI/ARCHITECTURE.md, Backend/src/domain/agents/agent_protocol.py
"""

import os
import pandas as pd
import json
import re
import shutil
from pathlib import Path
from datetime import datetime
from typing import Dict, Any, List, Optional, Tuple
import asyncio
import logging
import math
from functools import wraps

# Imports del dominio del proyecto
from src.domain.agents.agent_protocol import BaseAgent, AgentRequest, AgentResponse, IntentType
from src.core.exceptions import AgentError, ValidationError
# Authorization service import with fallback
try:
    from src.application.services.agent_authorization_service import authorize_agent_operation, AuthorizationError
except ImportError:
    # Fallback when authorization service is not available
    def authorize_agent_operation(agent_name: str, operation: str, path: str = None) -> bool:
        """Fallback authorization - always allow for now"""
        return True

    class AuthorizationError(Exception):
        pass

# Alias para compatibilidad
AgentResult = AgentResponse

logger = logging.getLogger(__name__)


class CapiDesktop(BaseAgent):
    """
    Capi Desktop

    Especialista en operaciones de archivos de oficina con enfoque en seguridad,
    precisión y preservación de datos. Como un desktop tradicional digital.
    """

    AGENT_NAME = "capi_desktop"
    VERSION = "1.0.0"
    DISPLAY_NAME = "Capi Desktop"

    SUPPORTED_INTENTS = [
        "leer_archivo_csv",
        "escribir_archivo_csv",
        "modificar_archivo_csv",
        "leer_archivo_excel",
        "escribir_archivo_excel",
        "modificar_archivo_excel",
        "leer_archivo_txt",
        "leer_archivo_word",
        "escribir_archivo_txt",
        "escribir_archivo_word",
        "modificar_archivo_word",
        "transformar_formato_archivo",
        "analizar_estructura_archivo",
        "validar_datos_archivo",
        "backup_archivo",
        "listar_archivos_desktop",
        "buscar_archivo_inteligente",
        "listar_archivos_similares"
    ]

    def __init__(self):
        super().__init__("capi_desktop")
        # Rutas para acceso al escritorio del usuario
        import os
        if os.path.exists("/app"):
            # En Docker priorizar el escritorio montado del usuario
            candidate_paths = [
                Path("/app/user_desktop"),
                Path("/app/ia_workspace/user_desktop"),
                Path("/app/ia_workspace/data/user_desktop"),
                Path("/app/ia_workspace/data")
            ]
            self.desktop_path = next((p for p in candidate_paths if p.exists()), Path("/app/ia_workspace/data"))
            self.backup_path = Path("/app/ia_workspace/data/backups")
        else:
            # En desarrollo local - usar el escritorio real del usuario
            self.desktop_path = Path("C:/Users/lucas/OneDrive/Desktop")
            self.backup_path = Path("C:/Users/lucas/OneDrive/Desktop/CAPI_Backups")
        handler_path = Path(__file__).resolve()
        self.workspace_root = handler_path.parents[2]
        self.data_root = self.workspace_root / "data"
        self.output_root = self.data_root / "agent-output" / self.AGENT_NAME

        self.data_root.mkdir(parents=True, exist_ok=True)
        self.output_root.mkdir(parents=True, exist_ok=True)

        self.max_file_size_mb = 50
        # Ampliar extensiones permitidas para incluir archivos sin extensión
        self.allowed_extensions = {'.csv', '.xlsx', '.xls', '.docx', '.doc', '.txt', ''}  # '' para archivos sin extensión

        # Crear carpeta de backups si no existe
        self.backup_path.mkdir(parents=True, exist_ok=True)

    @property
    def supported_intents(self) -> list[IntentType]:
        """Return list of intents this agent can handle."""
        return [IntentType.FILE_OPERATION]

    async def process(self, task) -> AgentResult:
        """Process method for BaseAgent compatibility"""
        # Convert task to AgentRequest format
        request = AgentRequest(
            intent=task.intent,
            query=task.query,
            parameters=task.metadata,
            user_id=task.user_id,
            session_id=task.session_id,
            context=task.context
        )
        return await self.handle(request)

    async def handle(self, request: AgentRequest) -> AgentResult:
        """
        Maneja solicitudes de operaciones de archivos con máxima seguridad
        """
        try:
            intent = request.intent
            params = request.parameters or {}

            logger.info(f"Capi Desktop procesando intent: {intent}")

            # Validaciones de seguridad
            await self._validate_request(request)

            # Router de intenciones
            if intent == "leer_archivo_csv":
                return await self._read_csv_file(params)
            elif intent == "escribir_archivo_csv":
                return await self._write_csv_file(params)
            elif intent == "modificar_archivo_csv":
                return await self._modify_csv_file(params)
            elif intent == "leer_archivo_excel":
                return await self._read_excel_file(params)
            elif intent == "escribir_archivo_excel":
                return await self._write_excel_file(params)
            elif intent == "modificar_archivo_excel":
                return await self._modify_excel_file(params)
            elif intent == "leer_archivo_txt":
                return await self._read_txt_file(params)
            elif intent == "leer_archivo_word":
                return await self._read_word_file(params)
            elif intent == "leer_archivo_pdf":
                return await self._read_pdf_file(params)
            elif intent == "transformar_formato_archivo":
                return await self._transform_file_format(params)
            elif intent == "analizar_estructura_archivo":
                return await self._analyze_file_structure(params)
            elif intent == "validar_datos_archivo":
                return await self._validate_file_data(params)
            elif intent == "backup_archivo":
                return await self._backup_file(params)
            elif intent == "listar_archivos_desktop":
                return await self._list_desktop_files(params)
            elif intent == "buscar_archivo_inteligente":
                return await self._intelligent_file_search(params)
            elif intent == "listar_archivos_similares":
                return await self._list_similar_files(params)
            else:
                raise AgentError(f"Intent no soportado: {intent}")

        except Exception as e:
            logger.error(f"Error en Capi Desktop: {str(e)}")
            return AgentResult(
                success=False,
                data={"error": str(e)},
                message=f"Error del Capi Desktop: {str(e)}",
                metadata={"agent": self.AGENT_NAME, "timestamp": datetime.now().isoformat()}
            )

    def _authorize_operation(self, operation: str, path: Optional[str] = None) -> bool:
        """Autorizar operación usando el sistema de privilegios"""
        try:
            return authorize_agent_operation(self.AGENT_NAME, operation, path)
        except AuthorizationError as e:
            logger.error(f"Error de autorización: {e}")
            return False

    async def _validate_request(self, request: AgentRequest) -> None:
        """Validaciones de seguridad estrictas"""
        params = request.parameters or {}

        # Validar archivo si se especifica
        if 'filename' in params:
            file_path = self._resolve_file_path(params['filename'])

            # Verificar extensión permitida
            if file_path.suffix.lower() not in self.allowed_extensions:
                raise ValidationError(f"Extensión no permitida: {file_path.suffix}")

            # Verificar tamaño si existe
            if file_path.exists():
                size_mb = file_path.stat().st_size / (1024 * 1024)
                if size_mb > self.max_file_size_mb:
                    raise ValidationError(f"Archivo demasiado grande: {size_mb:.1f}MB > {self.max_file_size_mb}MB")

    def _resolve_file_path(self, filename: str) -> Path:
        """Busca archivos inteligentemente en múltiples ubicaciones"""
        import os
        from pathlib import Path

        filename = (filename or '').strip()
        drive_fix = re.match(r'^([A-Za-z]):(?![\\/])', filename)
        if drive_fix:
            filename = f"{drive_fix.group(1)}:/" + filename[2:]

        if re.match(r'^[A-Za-z]:/[A-Za-z0-9]+', filename) and '/' not in filename.split(':', 1)[1]:
            filename = filename.replace('\\', '/').replace('\\', '/')

        tokens = ['Users', 'lucas', 'OneDrive', 'Desktop', 'Documents', 'Downloads']
        for token in tokens:
            filename = re.sub(rf'{token}(?![\/])', f'{token}/', filename)
        filename = re.sub(r'/+', '/', filename)

        # Lista de ubicaciones para buscar (en orden de prioridad)
        search_locations = []

        # Verificar si estamos en Docker o desarrollo local
        # Docker check más específico
        is_docker = os.path.exists("/app") and os.path.exists("/.dockerenv")
        if is_docker:
            # En Docker - incluir el escritorio montado
            search_locations = [
                Path("/app/user_desktop"),              # Escritorio del usuario montado (PRIORIDAD)
                Path("/app/ia_workspace/data"),
                Path("/app/data"),
                Path("/tmp")
            ]
        else:
            # En desarrollo local - buscar en ubicaciones comunes
            home = Path.home()
            search_locations = [
                Path("C:/Users/lucas/OneDrive/Desktop"),  # Escritorio OneDrive del usuario (PRIORIDAD)
                home / "Desktop",                    # Escritorio clásico del usuario
                home / "Documents",                  # Documentos
                home / "Downloads",                  # Descargas
                Path("C:/Users/lucas/Desktop"),      # Ruta específica del usuario
                Path("C:/Users/lucas/OneDrive/Desktop/CAPI"),  # Proyecto CAPI
                self.desktop_path,                   # Ubicación configurada
                Path.cwd(),                          # Directorio actual
            ]

        # Si es ruta absoluta, probar directamente primero
        if Path(filename).is_absolute():
            abs_path = Path(filename)
            if abs_path.exists():
                return abs_path
            # Si no existe, seguir buscando en las ubicaciones comunes

        # Buscar en todas las ubicaciones
        for location in search_locations:
            if not location.exists():
                continue

            # Probar nombre exacto
            candidate = location / filename
            if candidate.exists():
                logger.info(f'[DESKTOP] Archivo encontrado: {candidate}')
                return candidate

            # Probar con diferentes extensiones comunes
            if '.' not in filename:  # Si no tiene extensión
                for ext in ['.txt', '.xlsx', '.csv', '.docx']:
                    candidate_ext = location / f"{filename}{ext}"
                    if candidate_ext.exists():
                        logger.info(f'[DESKTOP] Archivo encontrado (con {ext}): {candidate_ext}')
                        return candidate_ext

            # Buscar con patrones similares (case insensitive)
            try:
                for existing_file in location.iterdir():
                    if existing_file.is_file():
                        # Comparación case-insensitive
                        if existing_file.name.lower() == filename.lower():
                            logger.info(f'[DESKTOP] Archivo encontrado (case-insensitive): {existing_file}')
                            return existing_file
                        if existing_file.name.lower() == f"{filename}.txt".lower():
                            logger.info(f'[DESKTOP] Archivo encontrado (case-insensitive con .txt): {existing_file}')
                            return existing_file
            except (PermissionError, OSError):
                continue

        # Si no se encuentra, devolver la ruta por defecto (que fallará)
        default_path = self.desktop_path / filename
        logger.warning(f'[DESKTOP] Archivo NO encontrado. Ubicaciones buscadas: {[str(loc) for loc in search_locations]} | solicitado={filename}')
        return default_path

    async def _create_backup(self, file_path: Path) -> Tuple[Path, Optional[str]]:
        """Crea backup del archivo antes de modificarlo"""
        if not file_path.exists():
            return file_path, None

        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        backup_name = f"{file_path.stem}_{timestamp}{file_path.suffix}"
        backup_file = self.backup_path / backup_name

        shutil.copy2(file_path, backup_file)
        archive_copy = self._archive_output_file(backup_file)
        logger.info(
            f"Backup creado: {backup_file} (copia archivada: {archive_copy})"
        )

        return backup_file, archive_copy

    def _archive_output_file(self, output_path: Path) -> Optional[str]:
        """Guarda una copia del archivo en la carpeta estandar de outputs del agente."""
        if output_path is None:
            return None

        output_path = Path(output_path)
        try:
            if not output_path.exists() or not output_path.is_file():
                return None

            now = datetime.now()
            archive_dir = self.output_root / f"{now.year:04d}" / f"{now.month:02d}"
            archive_dir.mkdir(parents=True, exist_ok=True)
            archive_name = f"{now.strftime('%Y%m%d_%H%M%S')}_{output_path.name}"
            archive_path = archive_dir / archive_name
            shutil.copy2(output_path, archive_path)
            logger.info(f"Copia archivada en {archive_path}")
            return str(archive_path)
        except Exception as exc:
            logger.warning(f"No se pudo archivar la salida {output_path}: {exc}")
            return None

    async def _read_csv_file(self, params: Dict[str, Any]) -> AgentResult:
        """Lee archivo CSV con encodings inteligentes"""
        filename = params.get('filename')
        if not filename:
            raise ValidationError("Parámetro 'filename' requerido")

        file_path = self._resolve_file_path(filename)

        if not file_path.exists():
            raise AgentError(f"Archivo no encontrado: {file_path}")

        # Intentar diferentes encodings
        encodings = ['utf-8', 'utf-8-sig', 'latin-1', 'cp1252', 'iso-8859-1']
        df = None
        encoding_used = None

        for encoding in encodings:
            try:
                # Auto-detectar delimitador usando engine python
                df = pd.read_csv(file_path, encoding=encoding, sep=None, engine='python')
                encoding_used = encoding
                break
            except UnicodeDecodeError:
                continue
            except Exception:
                # Intentar con coma por defecto si el autodetect falla por alguna razón
                try:
                    df = pd.read_csv(file_path, encoding=encoding)
                    encoding_used = encoding
                    break
                except Exception:
                    continue

        if df is None:
            raise AgentError("No se pudo leer el archivo con ningún encoding")

        # Análisis básico del archivo
        analysis = {
            "filename": filename,
            "rows": len(df),
            "columns": len(df.columns),
            "column_names": list(df.columns),
            "encoding_used": encoding_used,
            "memory_usage": df.memory_usage().sum(),
            "data_types": df.dtypes.to_dict()
        }

        # Retornar muestra de datos (primeras 10 filas)
        sample_data = df.head(10).to_dict('records')

        return AgentResult(
            success=True,
            data={
                "analysis": analysis,
                "sample_data": sample_data,
                "full_data_available": True
            },
            message=f"Archivo CSV leído: {len(df)} filas, {len(df.columns)} columnas",
            metadata={
                "agent": self.AGENT_NAME,
                "operation": "read_csv",
                "file_size_mb": round(file_path.stat().st_size / (1024 * 1024), 2)
            }
        )

    async def _write_csv_file(self, params: Dict[str, Any]) -> AgentResult:
        """Escribe archivo CSV con validaciones"""
        filename = params.get('filename')
        data = params.get('data')

        if not filename or not data:
            raise ValidationError("Parámetros 'filename' y 'data' requeridos")

        file_path = self._resolve_file_path(filename)

        # Crear backup si el archivo existe
        if file_path.exists():
            await self._create_backup(file_path)

        # Convertir datos a DataFrame
        if isinstance(data, list) and len(data) > 0:
            df = pd.DataFrame(data)
        elif isinstance(data, dict):
            df = pd.DataFrame([data])
        else:
            raise ValidationError("Formato de datos no válido")

        # Escribir archivo
        df.to_csv(file_path, index=False, encoding='utf-8')
        archive_path = self._archive_output_file(file_path)

        result_data = {
            "bytes_written": file_path.stat().st_size,
            "filename": filename,
            "rows_written": len(df),
            "columns_written": len(df.columns),
            "file_size_bytes": file_path.stat().st_size
        }
        if archive_path:
            result_data["archived_copy"] = archive_path

        return AgentResult(
            success=True,
            data=result_data,
            message=f"Archivo CSV creado: {filename} ({len(df)} filas)",
            metadata={"agent": self.AGENT_NAME, "operation": "write_csv"}
        )

    async def _list_desktop_files(self, params: Dict[str, Any]) -> AgentResult:
        """Lista archivos compatibles en el desktop y opcionalmente lee el primero"""
        pattern = params.get('pattern', '*')
        auto_read_first = params.get('auto_read_first', False)

        files = []
        file_paths = []

        logger.info("🔍 Iniciando búsqueda de archivos en escritorio...")

        # Buscar en múltiples ubicaciones, priorizando el escritorio real del usuario
        search_locations = [
            Path("/app/user_desktop"),                    # Escritorio montado en Docker (PRIORIDAD)
            Path("C:/Users/lucas/OneDrive/Desktop"),      # Escritorio principal del usuario (local)
            Path("C:/Users/lucas/Desktop"),               # Escritorio alternativo
            Path.home() / "Desktop",                      # Escritorio detectado automáticamente
            Path.home() / "OneDrive" / "Desktop",         # OneDrive Desktop
            self.desktop_path,                            # Configuración del agente
            Path("/app/ia_workspace/data"),               # Fallback en Docker
        ]

        logger.info(f"🔍 Buscando archivos en {len(search_locations)} ubicaciones...")
        for i, loc in enumerate(search_locations):
            logger.info(f"  [{i+1}] {loc} (existe: {loc.exists() if loc else False})")

        for location in search_locations:
            if not location.exists():
                continue

            # Buscar archivos con diferentes patrones
            for file_path in location.glob(pattern):
                if file_path.is_file() and file_path.suffix.lower() in self.allowed_extensions:
                    stat = file_path.stat()
                    files.append({
                        "name": file_path.name,
                        "size_mb": round(stat.st_size / (1024 * 1024), 3),
                        "modified": datetime.fromtimestamp(stat.st_mtime).isoformat(),
                        "extension": file_path.suffix.lower(),
                        "path": str(file_path)
                    })
                    file_paths.append(file_path)


        # Si se solicita auto-lectura y hay archivos, leer el primero disponible
        if auto_read_first and files:
            # Priorizar archivos de oficina por orden de frecuencia de uso
            office_extensions = ['.xlsx', '.xls', '.docx', '.doc', '.csv', '.txt', '.pdf']
            target_file = None

            # Buscar el primer archivo con extensión de oficina en orden de prioridad
            for ext in office_extensions:
                for file_path in file_paths:
                    if file_path.suffix.lower() == ext:
                        target_file = file_path
                        break
                if target_file:
                    break

            # Si no se encuentra archivo de oficina, usar el primero disponible
            if not target_file and file_paths:
                target_file = file_paths[0]

            if target_file:
                try:
                    # Leer el contenido del archivo según su extensión
                    if target_file.suffix.lower() in ['.xlsx', '.xls']:
                        content_result = await self._read_excel_file({"filename": target_file.name})
                    elif target_file.suffix.lower() == '.csv':
                        content_result = await self._read_csv_file({"filename": target_file.name})
                    elif target_file.suffix.lower() in ['.docx', '.doc']:
                        content_result = await self._read_word_file({"filename": target_file.name})
                    elif target_file.suffix.lower() == '.pdf':
                        content_result = await self._read_pdf_file({"filename": target_file.name})
                    else:
                        content_result = await self._read_txt_file({"filename": target_file.name})

                    if content_result.success:
                        return AgentResult(
                            success=True,
                            data={
                                "files": files,
                                "total_files": len(files),
                                "auto_read_file": target_file.name,
                                "content": content_result.data.get("content", ""),
                                "content_preview": content_result.data.get("content", "")[:200] + "..." if len(content_result.data.get("content", "")) > 200 else content_result.data.get("content", "")
                            },
                            message=f"📄 Archivo '{target_file.name}' encontrado y leído automáticamente.\n\n📝 Contenido:\n{content_result.data.get('content', '')}",
                            metadata={"agent": self.AGENT_NAME, "operation": "auto_read_file", "file_read": target_file.name}
                        )
                except Exception as e:
                    logger.error(f"Error leyendo archivo automáticamente: {e}")

        return AgentResult(
            success=True,
            data={"files": files, "total_files": len(files)},
            message=f"Encontrados {len(files)} archivos compatibles" + (f"\n\nArchivos encontrados:\n" + "\n".join([f"• {f['name']}" for f in files[:5]]) if files else ""),
            metadata={"agent": self.AGENT_NAME, "operation": "list_files"}
        )

    # Métodos placeholder para otras operaciones
    async def _modify_csv_file(self, params: Dict[str, Any]) -> AgentResult:
        """Modifica archivo CSV existente con un conjunto de operaciones seguras.

        Parámetros esperados:
        - filename: str (requerido)
        - operations: List[Dict[str, Any]] (requerido) con formas soportadas:
            {"type": "multiply_column", "column": "precio", "factor": 1.21}
            {"type": "add_constant", "column": "impuesto", "value": 5}
            {"type": "rename_columns", "mapping": {"old": "new"}}
            {"type": "drop_columns", "columns": ["col1", "col2"]}
            {"type": "fillna", "column": "monto", "value": 0}
            {"type": "filter_equals", "column": "estado", "value": "activo"}
        - output_filename: Optional[str] (por defecto sobre-escribe el original con backup)
        """

        filename = params.get('filename')
        operations = params.get('operations')
        output_filename = params.get('output_filename')

        if not filename or not operations:
            raise ValidationError("Parámetros 'filename' y 'operations' requeridos")

        src_path = self._resolve_file_path(filename)
        if not src_path.exists():
            raise AgentError(f"Archivo no encontrado: {src_path}")
        if src_path.suffix.lower() != '.csv':
            raise ValidationError("_modify_csv_file solo soporta .csv")

        # Leer CSV con heurísticas de encoding
        encodings = ['utf-8', 'utf-8-sig', 'latin-1', 'cp1252', 'iso-8859-1']
        last_err = None
        for encoding in encodings:
            try:
                df = pd.read_csv(src_path, encoding=encoding, sep=None, engine='python')
                encoding_used = encoding
                break
            except Exception as e:
                last_err = e
                df = None
        if df is None:
            raise AgentError(f"No se pudo leer el CSV: {last_err}")

        original_shape = df.shape
        steps: List[str] = []

        # Aplicar operaciones de forma segura
        for op in operations:
            if not isinstance(op, dict) or 'type' not in op:
                continue
            optype = str(op['type']).lower()

            if optype == 'multiply_column':
                col = op.get('column')
                factor = op.get('factor')
                if col in df.columns and isinstance(factor, (int, float)):
                    df[col] = pd.to_numeric(df[col], errors='coerce') * float(factor)
                    steps.append(f"multiply_column {col} x {factor}")
            elif optype == 'add_constant':
                col = op.get('column')
                value = op.get('value')
                if col:
                    if col not in df.columns:
                        df[col] = value
                    else:
                        # si es numérico, suma; si no, reemplaza NA
                        if pd.api.types.is_numeric_dtype(df[col]):
                            df[col] = pd.to_numeric(df[col], errors='coerce').fillna(0) + (value or 0)
                        else:
                            df[col] = df[col].fillna(value)
                    steps.append(f"add_constant {col} = {value}")
            elif optype == 'rename_columns':
                mapping = op.get('mapping') or {}
                if isinstance(mapping, dict) and mapping:
                    df = df.rename(columns=mapping)
                    steps.append(f"rename_columns {list(mapping.items())}")
            elif optype == 'drop_columns':
                cols = [c for c in (op.get('columns') or []) if c in df.columns]
                if cols:
                    df = df.drop(columns=cols)
                    steps.append(f"drop_columns {cols}")
            elif optype == 'fillna':
                col = op.get('column')
                value = op.get('value')
                if col in df.columns:
                    df[col] = df[col].fillna(value)
                    steps.append(f"fillna {col} -> {value}")
            elif optype == 'filter_equals':
                col = op.get('column')
                value = op.get('value')
                if col in df.columns:
                    prev_rows = len(df)
                    df = df[df[col] == value]
                    steps.append(f"filter_equals {col} == {value} (de {prev_rows} a {len(df)})")
            else:
                steps.append(f"operación no soportada: {optype}")

        # Decidir destino y backups
        if output_filename:
            dst_path = self._resolve_file_path(output_filename)
        else:
            dst_path = src_path
            # backup del original si vamos a sobrescribir
            await self._create_backup(src_path)

        # Escribir CSV con utf-8
        df.to_csv(dst_path, index=False, encoding='utf-8')

        result_data = {
            "source": str(src_path),
            "destination": str(dst_path),
            "original_shape": {
                "rows": int(original_shape[0]),
                "cols": int(original_shape[1])
            },
            "new_shape": {
                "rows": int(df.shape[0]),
                "cols": int(df.shape[1])
            },
            "steps": steps,
        }
        if archive_path:
            result_data["archived_copy"] = archive_path

        return AgentResult(
            success=True,
            data=result_data,
            message=f"CSV modificado y guardado en {dst_path.name}",
            metadata={"agent": self.AGENT_NAME, "operation": "modify_csv", "encoding": encoding_used}
        )

    async def _read_excel_file(self, params: Dict[str, Any]) -> AgentResult:
        """Lee archivo Excel (.xlsx/.xls) y devuelve análisis y muestra.

        Parámetros:
        - filename: str (requerido)
        - sheet_name: Optional[str|int] (por defecto primera hoja)
        - nrows_sample: int (por defecto 10)
        """
        filename = params.get('filename')
        if not filename:
            raise ValidationError("Parámetro 'filename' requerido")

        file_path = self._resolve_file_path(filename)
        if not file_path.exists():
            # PROFESSIONAL ERROR RECOVERY: Use intelligent file search instead of failing
            logger.info(f"File not found with exact name '{filename}', attempting intelligent search...")
            search_result = await self._intelligent_file_search({
                'original_filename': filename,
                'search_locations': ['desktop', 'documents', 'downloads']
            })

            # If intelligent search succeeded, return its result
            if search_result.success and search_result.data.get('auto_read_success'):
                logger.info(f"Professional recovery successful: {search_result.message}")
                return search_result

            # If search found suggestions but no auto-read, still show them
            if search_result.success and search_result.data.get('suggestions'):
                return search_result

            # Only raise error if intelligent search also failed
            raise AgentError(f"Archivo no encontrado: {file_path}")
        if file_path.suffix.lower() not in {'.xlsx', '.xls'}:
            raise ValidationError("Extensión no soportada para lectura de Excel")

        sheet_name = params.get('sheet_name', 0)
        nrows_sample = int(params.get('nrows_sample', 10))

        # Intentar con openpyxl si es xlsx
        read_kwargs = {"sheet_name": sheet_name}
        if file_path.suffix.lower() == '.xlsx':
            try:
                import openpyxl  # noqa: F401
                read_kwargs["engine"] = "openpyxl"
            except Exception:
                # Caer al engine por defecto; reportar en metadata
                pass

        df = pd.read_excel(file_path, **read_kwargs)

        analysis = {
            "filename": filename,
            "rows": int(len(df)),
            "columns": int(len(df.columns)),
            "column_names": list(map(str, df.columns)),
            "dtypes": {str(k): str(v) for k, v in df.dtypes.to_dict().items()},
        }

        sample_data = df.head(nrows_sample).to_dict('records')

        # Enhanced message with column content for better user experience
        column_info = f"Columnas: {', '.join(df.columns)}" if len(df.columns) > 0 else ""
        base_message = f"Excel leído: {len(df)} filas, {len(df.columns)} columnas"
        enhanced_message = f"{base_message}. {column_info}" if column_info else base_message

        response_data = {
            "analysis": analysis,
            "sample_data": sample_data,
            "columns": list(df.columns)
        }

        message_parts = [enhanced_message]

        if params.get('create_txt_copy'):
            try:
                txt_info = await self._create_txt_copy_from_dataframe(
                    df, file_path, params
                )
                response_data["generated_txt"] = txt_info
                message_parts.append(f"Copia txt: {txt_info['path']}")
            except Exception as copy_error:
                logger.warning(f"No se pudo generar copia txt: {copy_error}")
                response_data["generated_txt_error"] = str(copy_error)

        final_message = " | ".join(message_parts)

        return AgentResult(
            success=True,
            data=response_data,
            message=final_message,
            metadata={"agent": self.AGENT_NAME, "operation": "read_excel", "sheet": sheet_name}
        )



    async def _create_txt_copy_from_dataframe(self, df: pd.DataFrame, source_path: Path, params: Dict[str, Any]) -> Dict[str, Any]:
        """Genera un archivo de texto con el contenido resumido de un DataFrame."""
        txt_filename = params.get('txt_filename') or f"{source_path.stem}_resumen"
        txt_extension = params.get('txt_extension', '.txt') or '.txt'

        sanitized_name = self._sanitize_output_name(txt_filename)
        target_path = self.desktop_path / sanitized_name

        if txt_extension:
            if not txt_extension.startswith('.'):
                txt_extension = f'.{txt_extension}'
            target_path = target_path.with_suffix(txt_extension)

        target_path.parent.mkdir(parents=True, exist_ok=True)

        if target_path.exists():
            await self._create_backup(target_path)

        content = self._format_excel_summary(df, source_path)
        target_path.write_text(content, encoding='utf-8')
        archive_copy = self._archive_output_file(target_path)

        result = {
            "path": str(target_path),
            "preview": content[:2000]
        }
        if archive_copy:
            result["archived_copy"] = archive_copy

        return result

    def _format_excel_summary(self, df: pd.DataFrame, source_path: Path, max_rows: int = 20) -> str:
        """Construye un resumen de texto a partir de un DataFrame."""
        header_lines = [
            f"Archivo: {source_path.name}",
            f"Filas: {len(df)} | Columnas: {len(df.columns)}"
        ]
        if len(df.columns) > 0:
            header_lines.append("Columnas: " + ", ".join(map(str, df.columns)))
        header_lines.append('-' * 80)

        preview_rows = min(max_rows, len(df))
        if preview_rows > 0:
            preview_df = df.head(preview_rows)
            preview_str = preview_df.to_string(index=False)
        else:
            preview_str = "[SIN DATOS EN LA HOJA]"

        return "\n".join(header_lines + [preview_str])

    def _sanitize_output_name(self, name: str) -> str:
        """Normaliza un nombre de archivo para salidas en el escritorio."""
        cleaned = re.sub(r'[^A-Za-z0-9._-]+', '_', (name or '').strip())
        if not cleaned:
            cleaned = f"{self.AGENT_NAME}_output"
        return cleaned

    async def _read_txt_file(self, params: Dict[str, Any]) -> AgentResult:
        """Lee archivo de texto (.txt) y devuelve su contenido.

        Parámetros:
        - filename: str (requerido)
        - encoding: str (opcional, por defecto 'utf-8')
        - max_lines: int (opcional, máximo de líneas a leer)
        """
        filename = params.get('filename')
        if not filename:
            raise ValidationError("Parámetro 'filename' requerido")

        file_path = self._resolve_file_path(filename)

        # AUTORIZACIÓN: Verificar permisos antes de acceder al archivo
        operation = "read_file"
        if str(file_path).startswith(str(Path.home())):
            operation = "read_desktop"
        elif not str(file_path).startswith(str(self.desktop_path)):
            operation = "access_external_path"

        if not self._authorize_operation(operation, str(file_path)):
            raise AuthorizationError(f"Acceso denegado al archivo: {file_path}")

        if not file_path.exists():
            raise AgentError(f"Archivo no encontrado: {file_path}")
        ext = file_path.suffix.lower()
        if ext not in {'.txt', ''}:
            redirect_params = dict(params)
            redirect_params['filename'] = file_path.name
            if ext in {'.xlsx', '.xls'}:
                return await self._read_excel_file(redirect_params)
            if ext == '.csv':
                return await self._read_csv_file(redirect_params)
            if ext in {'.docx', '.doc'}:
                return await self._read_word_file(redirect_params)
            raise ValidationError('Extensión no soportada para lectura de texto')


        encoding = params.get('encoding', 'utf-8')
        max_lines = params.get('max_lines', 100)  # Limit for safety

        try:
            with open(file_path, 'r', encoding=encoding, errors='replace') as f:
                lines = []
                for i, line in enumerate(f):
                    if i >= max_lines:
                        break
                    lines.append(line.rstrip('\n\r'))

                content = '\n'.join(lines)
                total_lines = i + 1

            # Basic analysis
            analysis = {
                "filename": filename,
                "total_lines": total_lines,
                "content_length": len(content),
                "encoding": encoding,
                "file_size_bytes": file_path.stat().st_size,
                "truncated": total_lines > max_lines
            }

            return AgentResult(
                success=True,
                data={"content": content, "analysis": analysis},
                message=f"Archivo de texto leído: {total_lines} líneas, {len(content)} caracteres",
                metadata={"agent": self.AGENT_NAME, "operation": "read_txt", "encoding": encoding}
            )

        except UnicodeDecodeError:
            # Try with different encodings
            for alt_encoding in ['latin-1', 'cp1252']:
                try:
                    with open(file_path, 'r', encoding=alt_encoding, errors='replace') as f:
                        content = f.read(max_lines * 100)  # Reasonable limit

                    analysis = {
                        "filename": filename,
                        "content_length": len(content),
                        "encoding": alt_encoding,
                        "file_size_bytes": file_path.stat().st_size,
                        "encoding_fallback": True
                    }

                    return AgentResult(
                        success=True,
                        data={"content": content, "analysis": analysis},
                        message=f"Archivo leído con encoding {alt_encoding}: {len(content)} caracteres",
                        metadata={"agent": self.AGENT_NAME, "operation": "read_txt", "encoding": alt_encoding}
                    )
                except:
                    continue

            raise AgentError(f"No se pudo leer el archivo con ningún encoding soportado")

    async def _write_txt_file(self, params: Dict[str, Any]) -> AgentResult:
        """Escribe un archivo de texto en el escritorio del usuario."""
        filename = params.get("filename")
        if not filename:
            raise ValidationError("Parámetro 'filename' requerido")

        content = params.get("content") or params.get("text")
        data = params.get("data")
        if content is None and data is not None:
            if isinstance(data, dict):
                content = json.dumps(data, ensure_ascii=False, indent=2)
            elif isinstance(data, (list, tuple)) and data:
                first = data[0]
                if isinstance(first, dict):
                    content = json.dumps(first, ensure_ascii=False, indent=2)
                else:
                    content = "\n".join(str(item) for item in data)
            else:
                content = str(data)
        if content is None:
            raise ValidationError("Parámetro 'content' o 'data' requerido para escribir el archivo")

        file_path = self._resolve_file_path(filename)
        if file_path.suffix.lower() != '.txt':
            file_path = file_path.with_suffix('.txt')
        file_path.parent.mkdir(parents=True, exist_ok=True)

        if file_path.exists():
            await self._create_backup(file_path)

        file_path.write_text(str(content), encoding='utf-8')
        archive_path = self._archive_output_file(file_path)

        result_data = {
            "filename": str(file_path)
        }
        if archive_path:
            result_data["archived_copy"] = archive_path

        return AgentResult(
            success=True,
            data=result_data,
            message=f"Archivo de texto escrito en {file_path}",
            metadata={"agent": self.AGENT_NAME, "operation": "write_txt"}
        )

    async def _write_excel_file(self, params: Dict[str, Any]) -> AgentResult:
        """Escribe archivo Excel a partir de 'data' o de un CSV existente.

        Parámetros:
        - filename: str (requerido, destino .xlsx)
        - data: List[Dict] | Dict (opcional)
        - from_csv: str (opcional, nombre de CSV en Desktop)
        - sheet_name: str (opcional, por defecto 'Sheet1')
        """
        filename = params.get('filename')
        data = params.get('data')
        from_csv = params.get('from_csv')
        sheet_name = params.get('sheet_name', 'Sheet1')

        if not filename:
            raise ValidationError("Parámetro 'filename' requerido")
        dst_path = self._resolve_file_path(filename)
        if dst_path.suffix.lower() != '.xlsx':
            raise ValidationError("El destino debe ser un archivo .xlsx")

        # Crear DF
        df = None
        if data is not None:
            if isinstance(data, list) and data:
                df = pd.DataFrame(data)
            elif isinstance(data, dict):
                df = pd.DataFrame([data])
            else:
                raise ValidationError("Formato de 'data' no válido")
        elif from_csv:
            csv_path = self._resolve_file_path(from_csv)
            if not csv_path.exists():
                raise AgentError(f"CSV de origen no encontrado: {csv_path}")
            df = pd.read_csv(csv_path, sep=None, engine='python')
        else:
            raise ValidationError("Debe proporcionar 'data' o 'from_csv'")

        # Escribir Excel (preferir openpyxl si está disponible)
        to_excel_kwargs = {"index": False, "sheet_name": sheet_name}
        try:
            import openpyxl  # noqa: F401
            to_excel_kwargs["engine"] = "openpyxl"
        except Exception:
            pass

        # Backup si sobrescribe
        if dst_path.exists():
            await self._create_backup(dst_path)

        with pd.ExcelWriter(dst_path, **({"engine": to_excel_kwargs.get("engine")} if to_excel_kwargs.get("engine") else {})) as writer:
            df.to_excel(writer, sheet_name=sheet_name, index=False)
        archive_path = self._archive_output_file(dst_path)

        result_data = {
            "filename": str(dst_path),
            "rows_written": int(len(df)),
            "columns_written": int(len(df.columns))
        }
        if archive_path:
            result_data["archived_copy"] = archive_path

        return AgentResult(
            success=True,
            data=result_data,
            message=f"Excel escrito: {dst_path.name} ({len(df)} filas)",
            metadata={"agent": self.AGENT_NAME, "operation": "write_excel"}
        )

    async def _modify_excel_file(self, params: Dict[str, Any]) -> AgentResult:
        """Modifica archivo Excel aplicando operaciones por hoja (soporte básico).

        Parámetros:
        - filename: str (requerido)
        - sheet_name: str|int (opcional, por defecto 0)
        - operations: List[Dict] (mismas que CSV cuando aplican)
        - output_filename: Optional[str]
        """
        filename = params.get('filename')
        operations = params.get('operations') or []
        sheet_name = params.get('sheet_name', 0)
        output_filename = params.get('output_filename')

        if not filename:
            raise ValidationError("Parámetro 'filename' requerido")
        src_path = self._resolve_file_path(filename)
        if not src_path.exists():
            raise AgentError(f"Archivo no encontrado: {src_path}")
        if src_path.suffix.lower() not in {'.xlsx', '.xls'}:
            raise ValidationError("Solo se soportan .xlsx/.xls")

        # Leer hoja solicitada
        df = pd.read_excel(src_path, sheet_name=sheet_name)
        csv_like_params = {"filename": "ignored.csv", "operations": operations, "output_filename": None}

        # Reutilizar lógica de operaciones sobre DataFrame:
        # Implementamos una mini-función local que comparte con CSV
        def apply_ops(df_in: pd.DataFrame, ops: List[Dict[str, Any]]) -> Tuple[pd.DataFrame, List[str]]:
            steps: List[str] = []
            df_local = df_in.copy()
            for op in ops:
                optype = (op.get('type') or '').lower()
                if optype == 'multiply_column' and op.get('column') in df_local.columns:
                    df_local[op['column']] = pd.to_numeric(df_local[op['column']], errors='coerce') * float(op.get('factor', 1))
                    steps.append(f"multiply_column {op['column']} x {op.get('factor', 1)}")
                elif optype == 'add_constant' and op.get('column'):
                    col = op['column']
                    val = op.get('value')
                    if col not in df_local.columns:
                        df_local[col] = val
                    else:
                        if pd.api.types.is_numeric_dtype(df_local[col]):
                            df_local[col] = pd.to_numeric(df_local[col], errors='coerce').fillna(0) + (val or 0)
                        else:
                            df_local[col] = df_local[col].fillna(val)
                    steps.append(f"add_constant {col} = {val}")
                elif optype == 'rename_columns':
                    mapping = op.get('mapping') or {}
                    df_local = df_local.rename(columns=mapping)
                    steps.append("rename_columns")
                elif optype == 'drop_columns':
                    cols = [c for c in (op.get('columns') or []) if c in df_local.columns]
                    if cols:
                        df_local = df_local.drop(columns=cols)
                        steps.append(f"drop_columns {cols}")
                elif optype == 'fillna' and op.get('column') in df_local.columns:
                    df_local[op['column']] = df_local[op['column']].fillna(op.get('value'))
                    steps.append(f"fillna {op['column']}")
                elif optype == 'filter_equals' and op.get('column') in df_local.columns:
                    df_local = df_local[df_local[op['column']] == op.get('value')]
                    steps.append(f"filter_equals {op['column']} == {op.get('value')}")
            return df_local, steps

        original_shape = df.shape
        df_new, steps = apply_ops(df, operations)

        # Seleccionar destino y backup
        if output_filename:
            dst_path = self._resolve_file_path(output_filename)
        else:
            dst_path = src_path
            await self._create_backup(src_path)

        # Escribir a Excel
        to_excel_kwargs = {"index": False}
        try:
            import openpyxl  # noqa: F401
            to_excel_kwargs["engine"] = "openpyxl"
        except Exception:
            pass

        with pd.ExcelWriter(dst_path, **({"engine": to_excel_kwargs.get("engine")} if to_excel_kwargs.get("engine") else {})) as writer:
            df_new.to_excel(writer, sheet_name=sheet_name if isinstance(sheet_name, str) else 'Sheet1', index=False)
        archive_path = self._archive_output_file(dst_path)

        result_data = {
            "source": str(src_path),
            "destination": str(dst_path),
            "original_shape": {
                "rows": int(original_shape[0]),
                "cols": int(original_shape[1])
            },
            "new_shape": {
                "rows": int(df_new.shape[0]),
                "cols": int(df_new.shape[1])
            },
            "steps": steps,
        }
        if archive_path:
            result_data["archived_copy"] = archive_path

        return AgentResult(
            success=True,
            data=result_data,
            message=f"Excel modificado y guardado en {dst_path.name}",
            metadata={"agent": self.AGENT_NAME, "operation": "modify_excel"}
        )

    async def _transform_file_format(self, params: Dict[str, Any]) -> AgentResult:
        """Transforma entre formatos soportados, preservando contenido tabular.

        Soporta: CSV <-> XLSX, TXT (delimitado) -> CSV

        Parámetros:
        - filename: str (requerido) archivo origen en Desktop
        - output_filename: str (requerido) nombre destino en Desktop
        - delimiter: Optional[str] para TXT/CSV
        - sheet_name: Optional[str] para Excel
        """
        src = params.get('filename')
        dst = params.get('output_filename')
        delimiter = params.get('delimiter')
        sheet_name = params.get('sheet_name', 'Sheet1')

        if not src or not dst:
            raise ValidationError("'filename' y 'output_filename' son requeridos")

        src_path = self._resolve_file_path(src)
        dst_path = self._resolve_file_path(dst)
        if not src_path.exists():
            raise AgentError(f"Archivo de origen no encontrado: {src_path}")

        src_ext = src_path.suffix.lower()
        dst_ext = dst_path.suffix.lower()

        if src_ext not in self.allowed_extensions or dst_ext not in self.allowed_extensions:
            raise ValidationError("Extensión no soportada para transformación")

        # Leer origen
        if src_ext == '.csv' or (src_ext == '.txt'):
            read_kwargs = {}
            if delimiter:
                read_kwargs['sep'] = delimiter
            else:
                read_kwargs['sep'] = None
                read_kwargs['engine'] = 'python'
            df = pd.read_csv(src_path, **read_kwargs)
        elif src_ext in {'.xlsx', '.xls'}:
            df = pd.read_excel(src_path)
        else:
            raise ValidationError("Transformación solo soporta CSV/XLSX/TXT")

        # Escribir destino
        if dst_ext == '.csv' or dst_ext == '.txt':
            if dst_path.exists():
                await self._create_backup(dst_path)
            df.to_csv(dst_path, index=False, encoding='utf-8')
        elif dst_ext == '.xlsx':
            if dst_path.exists():
                await self._create_backup(dst_path)
            to_excel_kwargs = {"index": False}
            try:
                import openpyxl  # noqa: F401
                to_excel_kwargs["engine"] = "openpyxl"
            except Exception:
                pass
            with pd.ExcelWriter(dst_path, **({"engine": to_excel_kwargs.get("engine")} if to_excel_kwargs.get("engine") else {})) as writer:
                df.to_excel(writer, sheet_name=sheet_name, index=False)
        else:
            raise ValidationError("Extensión de destino no soportada")

        archive_path = self._archive_output_file(dst_path)
        result_data = {
            "source": str(src_path),
            "destination": str(dst_path),
            "rows": int(len(df)),
            "columns": int(len(df.columns))
        }
        if archive_path:
            result_data["archived_copy"] = archive_path

        return AgentResult(
            success=True,
            data=result_data,
            message=f"Transformación completada: {src_path.name} -> {dst_path.name}",
            metadata={"agent": self.AGENT_NAME, "operation": "transform"}
        )

    async def _analyze_file_structure(self, params: Dict[str, Any]) -> AgentResult:
        """Analiza estructura y calidad de un archivo tabular (CSV/XLSX/TXT)."""
        filename = params.get('filename')
        if not filename:
            raise ValidationError("Parámetro 'filename' requerido")

        file_path = self._resolve_file_path(filename)
        if not file_path.exists():
            raise AgentError(f"Archivo no encontrado: {file_path}")

        ext = file_path.suffix.lower()
        if ext in {'.csv', '.txt'}:
            try:
                df = pd.read_csv(file_path, sep=None, engine='python')
            except Exception:
                df = pd.read_csv(file_path)
        elif ext in {'.xlsx', '.xls'}:
            df = pd.read_excel(file_path)
        else:
            raise ValidationError("Solo se analiza CSV/XLSX/TXT")

        null_counts = df.isna().sum().to_dict()
        duplicate_rows = int(df.duplicated().sum())
        memory_usage = int(df.memory_usage(deep=True).sum())
        sample = df.head(10).to_dict('records')

        report = {
            "filename": file_path.name,
            "shape": {"rows": int(len(df)), "columns": int(len(df.columns))},
            "columns": list(map(str, df.columns)),
            "dtypes": {str(k): str(v) for k, v in df.dtypes.to_dict().items()},
            "null_counts": {str(k): int(v) for k, v in null_counts.items()},
            "duplicate_rows": duplicate_rows,
            "memory_usage_bytes": memory_usage,
            "head_sample": sample,
        }

        return AgentResult(
            success=True,
            data={"report": report},
            message=f"Análisis de estructura completado: {len(df)} filas, {len(df.columns)} columnas",
            metadata={"agent": self.AGENT_NAME, "operation": "analyze_structure"}
        )

    async def _validate_file_data(self, params: Dict[str, Any]) -> AgentResult:
        """Valida integridad de datos del archivo basado en reglas proporcionadas.

        Parámetros:
        - filename: str (requerido)
        - required_columns: List[str]
        - unique_columns: List[str]
        - no_null_columns: List[str]
        - column_types: Dict[str, str] (int|float|str|datetime)
        - ranges: Dict[str, Dict[str, Any]] ej: {"precio": {"min": 0, "max": 1000}}
        """
        filename = params.get('filename')
        if not filename:
            raise ValidationError("Parámetro 'filename' requerido")

        file_path = self._resolve_file_path(filename)
        if not file_path.exists():
            raise AgentError(f"Archivo no encontrado: {file_path}")

        ext = file_path.suffix.lower()
        if ext in {'.csv', '.txt'}:
            df = pd.read_csv(file_path, sep=None, engine='python')
        elif ext in {'.xlsx', '.xls'}:
            df = pd.read_excel(file_path)
        else:
            raise ValidationError("Solo se valida CSV/XLSX/TXT")

        issues: List[str] = []
        checks: Dict[str, Any] = {}

        # Columnas requeridas
        required_columns = params.get('required_columns', [])
        missing = [c for c in required_columns if c not in df.columns]
        checks['missing_required_columns'] = missing
        if missing:
            issues.append(f"Faltan columnas requeridas: {missing}")

        # Columnas únicas
        unique_columns = params.get('unique_columns', [])
        unique_violations = {}
        for c in unique_columns:
            if c in df.columns:
                dup = int(df[c].duplicated().sum())
                if dup > 0:
                    unique_violations[c] = dup
        if unique_violations:
            issues.append(f"Violaciones de unicidad: {unique_violations}")
        checks['unique_violations'] = unique_violations

        # No null en columnas específicas
        no_null_columns = params.get('no_null_columns', [])
        null_violations = {}
        for c in no_null_columns:
            if c in df.columns:
                cnt = int(df[c].isna().sum())
                if cnt > 0:
                    null_violations[c] = cnt
        if null_violations:
            issues.append(f"Valores nulos en columnas prohibidas: {null_violations}")
        checks['null_violations'] = null_violations

        # Tipos de columnas
        type_map = {"int": 'int64', "float": 'float64', "str": 'object', "datetime": 'datetime64[ns]'}
        column_types = params.get('column_types', {})
        type_mismatches = {}
        for col, t in column_types.items():
            if col in df.columns:
                desired = str(t).lower()
                try:
                    if desired == 'int':
                        df[col] = pd.to_numeric(df[col], errors='coerce').astype('Int64')
                    elif desired == 'float':
                        df[col] = pd.to_numeric(df[col], errors='coerce')
                    elif desired == 'datetime':
                        df[col] = pd.to_datetime(df[col], errors='coerce')
                    elif desired == 'str':
                        df[col] = df[col].astype('string')
                except Exception:
                    type_mismatches[col] = desired
        if type_mismatches:
            issues.append(f"Columnas con tipos no convertibles: {type_mismatches}")
        checks['type_conversions_applied'] = list(column_types.keys())

        # Rangos
        ranges = params.get('ranges', {})
        range_violations = {}
        for col, spec in ranges.items():
            if col in df.columns:
                s = pd.to_numeric(df[col], errors='coerce')
                vmin = spec.get('min')
                vmax = spec.get('max')
                violations = 0
                if vmin is not None:
                    violations += int((s < vmin).sum())
                if vmax is not None:
                    violations += int((s > vmax).sum())
                if violations:
                    range_violations[col] = violations
        if range_violations:
            issues.append(f"Valores fuera de rango: {range_violations}")
        checks['range_violations'] = range_violations

        valid = len(issues) == 0

        return AgentResult(
            success=valid,
            data={"checks": checks, "total_rows": int(len(df))},
            message=("Datos válidos" if valid else "Se encontraron problemas de validación"),
            metadata={"agent": self.AGENT_NAME, "operation": "validate_data", "issues": issues}
        )

    async def _backup_file(self, params: Dict[str, Any]) -> AgentResult:
        """Crea backup manual de archivo"""
        filename = params.get('filename')
        if not filename:
            raise ValidationError("Parámetro 'filename' requerido")

        file_path = self._resolve_file_path(filename)
        backup_path, archive_copy = await self._create_backup(file_path)
        result_data = {
            "original": str(file_path),
            "backup": str(backup_path)
        }
        if archive_copy:
            result_data["archived_copy"] = archive_copy


        return AgentResult(
            success=True,
            data=result_data,
            message=f"Backup creado: {backup_path.name}",
            metadata={"agent": self.AGENT_NAME, "operation": "backup"}
        )

    async def _read_word_file(self, params: Dict[str, Any]) -> AgentResult:
        """Lee contenido de archivos Word (.docx/.doc)"""
        filename = params.get('filename')
        if not filename:
            raise ValidationError("Parámetro 'filename' requerido")

        file_path = self._resolve_file_path(filename)
        if not file_path.exists():
            raise AgentError(f"Archivo no encontrado: {file_path}")

        try:
            # Try python-docx for .docx files
            if file_path.suffix.lower() == '.docx':
                try:
                    import docx
                    doc = docx.Document(file_path)
                    content_parts = []
                    for paragraph in doc.paragraphs:
                        if paragraph.text.strip():
                            content_parts.append(paragraph.text.strip())
                    content = '\n'.join(content_parts)

                    analysis = {
                        "filename": filename,
                        "content_length": len(content),
                        "paragraphs": len(doc.paragraphs),
                        "file_size_bytes": file_path.stat().st_size
                    }

                    return AgentResult(
                        success=True,
                        data={"content": content, "analysis": analysis},
                        message=f"Archivo Word leído: {len(doc.paragraphs)} párrafos, {len(content)} caracteres",
                        metadata={"agent": self.AGENT_NAME, "operation": "read_word"}
                    )
                except ImportError:
                    # Fallback: treat as plain text if docx library not available
                    pass

            # Fallback: try as plain text
            return await self._read_txt_file({"filename": filename})

        except Exception as e:
            raise AgentError(f"Error leyendo archivo Word: {str(e)}")

    async def _read_pdf_file(self, params: Dict[str, Any]) -> AgentResult:
        """Lee contenido de archivos PDF (básico)"""
        filename = params.get('filename')
        if not filename:
            raise ValidationError("Parámetro 'filename' requerido")

        file_path = self._resolve_file_path(filename)
        if not file_path.exists():
            raise AgentError(f"Archivo no encontrado: {file_path}")

        try:
            # Try PyPDF2 for PDF reading
            try:
                import PyPDF2
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    content_parts = []
                    for page in pdf_reader.pages[:10]:  # Limit to first 10 pages
                        text = page.extract_text()
                        if text.strip():
                            content_parts.append(text.strip())
                    content = '\n'.join(content_parts)

                    analysis = {
                        "filename": filename,
                        "content_length": len(content),
                        "total_pages": len(pdf_reader.pages),
                        "pages_read": min(10, len(pdf_reader.pages)),
                        "file_size_bytes": file_path.stat().st_size
                    }

                    return AgentResult(
                        success=True,
                        data={"content": content, "analysis": analysis},
                        message=f"Archivo PDF leído: {analysis['pages_read']} de {analysis['total_pages']} páginas, {len(content)} caracteres",
                        metadata={"agent": self.AGENT_NAME, "operation": "read_pdf"}
                    )
            except ImportError:
                # PDF library not available
                raise AgentError("Librería PyPDF2 no disponible. No se pueden leer archivos PDF.")

        except Exception as e:
            raise AgentError(f"Error leyendo archivo PDF: {str(e)}")

    async def _intelligent_file_search(self, params: Dict[str, Any]) -> AgentResult:
        """
        Búsqueda inteligente de archivos con fuzzy matching y sugerencias.
        Profesionalización del orquestador para manejo de errores.
        """
        original_filename = params.get('original_filename', params.get('filename', ''))
        search_locations = params.get('search_locations', ['desktop', 'documents', 'downloads'])

        if not original_filename:
            raise ValidationError("Parámetro 'original_filename' requerido para búsqueda inteligente")

        found_files = []

        # Buscar en todas las ubicaciones configuradas
        for location_name in search_locations:
            location_paths = self._get_location_paths(location_name)

            for location in location_paths:
                if not location.exists():
                    continue

                try:
                    for file_path in location.rglob("*"):
                        if file_path.is_file():
                            similarity = self._calculate_filename_similarity(
                                original_filename, file_path.name
                            )
                            if similarity > 0.4:  # Threshold para considerar similitud
                                found_files.append({
                                    'path': str(file_path),
                                    'name': file_path.name,
                                    'similarity': similarity,
                                    'location': location_name
                                })
                except (PermissionError, OSError):
                    continue

        # Ordenar por similitud descendente
        found_files.sort(key=lambda x: x['similarity'], reverse=True)

        if found_files:
            # Intentar leer el archivo más similar automáticamente
            best_match = found_files[0]
            if best_match['similarity'] > 0.8:  # Alta confianza
                try:
                    # Intentar leer el archivo automáticamente
                    file_path = Path(best_match['path'])
                    if file_path.suffix.lower() in {'.xlsx', '.xls'}:
                        # Read file directly bypassing the recursive call
                        try:
                            import pandas as pd
                            df = pd.read_excel(file_path, sheet_name=0)

                            analysis = {
                                "filename": file_path.name,
                                "rows": int(len(df)),
                                "columns": int(len(df.columns)),
                                "column_names": list(map(str, df.columns)),
                                "dtypes": {str(k): str(v) for k, v in df.dtypes.to_dict().items()},
                            }

                            sample_data = df.head(10).to_dict('records')

                            return AgentResult(
                                success=True,
                                data={
                                    "auto_read_success": True,
                                    "analysis": analysis,
                                    "sample_data": sample_data,
                                    "file_found": best_match
                                },
                                message=f"✅ Encontré y leí automáticamente: '{best_match['name']}' (similitud: {best_match['similarity']:.2f})\n\nExcel leído: {len(df)} filas, {len(df.columns)} columnas\nColumnas: {', '.join(df.columns[:5])}{'...' if len(df.columns) > 5 else ''}\n\nPrimeras filas:\n" + "\n".join([f"{i+1}. {' | '.join([str(v)[:50] for v in row.values()][:3])}{'...' if len(row) > 3 else ''}" for i, row in enumerate(sample_data[:5])]),
                                metadata={"agent": self.AGENT_NAME, "operation": "intelligent_read", "auto_recovery": True}
                            )
                        except Exception as read_error:
                            logger.warning(f"Failed to read found file: {read_error}")
                            pass  # Continue to suggestions
                except Exception:
                    pass  # Si falla la lectura automática, continúa con sugerencias

            # Mostrar sugerencias al usuario
            suggestions = found_files[:5]  # Top 5 sugerencias
            suggestion_text = "\n".join([
                f"• {file['name']} (similitud: {file['similarity']:.2f}) en {file['location']}"
                for file in suggestions
            ])

            return AgentResult(
                success=True,
                data={
                    "suggestions": suggestions,
                    "original_filename": original_filename,
                    "search_performed": True
                },
                message=f"No encontré exactamente '{original_filename}', pero encontré archivos similares:\n\n{suggestion_text}\n\n💡 Especifica el nombre completo del archivo que quieres leer.",
                metadata={'agent': getattr(self, 'AGENT_NAME', 'capi_desktop')}
            )
        else:
            return AgentResult(
                success=False,
                data={"original_filename": original_filename, "search_performed": True},
                message=f"❌ No encontré archivos similares a '{original_filename}' en las ubicaciones de búsqueda.",
                metadata={'agent': getattr(self, 'AGENT_NAME', 'capi_desktop')}
            )

    async def _list_similar_files(self, params: Dict[str, Any]) -> AgentResult:
        """Lista archivos similares al patrón especificado"""
        pattern = params.get('pattern', '*')
        suggestion_mode = params.get('suggestion_mode', False)

        files_found = []
        search_locations = self._get_all_search_locations()

        for location in search_locations:
            if not location.exists():
                continue

            try:
                for file_path in location.glob(pattern):
                    if file_path.is_file():
                        files_found.append({
                            'name': file_path.name,
                            'path': str(file_path),
                            'size': file_path.stat().st_size,
                            'location': str(location)
                        })
            except (PermissionError, OSError):
                continue

        if files_found:
            files_text = "\n".join([
                f"• {file['name']} ({file['size']} bytes)"
                for file in files_found[:10]  # Limitar a 10 resultados
            ])

            if suggestion_mode:
                message = f"Archivos que podrían interesarte:\n\n{files_text}"
            else:
                message = f"Archivos encontrados con patrón '{pattern}':\n\n{files_text}"

            return AgentResult(
                success=True,
                data={"files": files_found, "pattern": pattern},
                message=message,
                metadata={'agent': getattr(self, 'AGENT_NAME', 'capi_desktop')}
            )
        else:
            return AgentResult(
                success=False,
                data={"pattern": pattern},
                message=f"No se encontraron archivos con el patrón '{pattern}'",
                metadata={'agent': getattr(self, 'AGENT_NAME', 'capi_desktop')}
            )

    def _get_location_paths(self, location_name: str) -> List[Path]:
        """Obtiene las rutas físicas para una ubicación lógica"""
        location_map = {
            'desktop': [
                getattr(self, 'desktop_path', None),
                Path('/app/user_desktop'),
                Path('C:/Users/lucas/OneDrive/Desktop'),
                Path.home() / 'Desktop',
                Path.home() / 'OneDrive' / 'Desktop'
            ],
            'documents': [
                Path.home() / 'Documents',
                Path.home() / 'OneDrive' / 'Documents'
            ],
            'downloads': [
                Path.home() / 'Downloads'
            ]
        }
        candidates = location_map.get(location_name, [])
        resolved_paths = []
        for candidate in candidates:
            if candidate is None:
                continue
            try:
                path_obj = Path(candidate)
            except TypeError:
                continue
            if path_obj not in resolved_paths:
                resolved_paths.append(path_obj)
        return resolved_paths

    def _get_all_search_locations(self) -> List[Path]:
        """Obtiene todas las ubicaciones de búsqueda"""
        locations = []
        for location_name in ['desktop', 'documents', 'downloads']:
            locations.extend(self._get_location_paths(location_name))
        return locations

    def _calculate_filename_similarity(self, name1: str, name2: str) -> float:
        """Calcula similitud entre nombres de archivo usando algoritmo simple"""
        name1_clean = name1.lower().replace(' ', '').replace('_', '').replace('-', '')
        name2_clean = name2.lower().replace(' ', '').replace('_', '').replace('-', '')

        # Remover extensiones para comparación
        name1_base = name1_clean.split('.')[0]
        name2_base = name2_clean.split('.')[0]

        # Similitud exacta
        if name1_base == name2_base:
            return 1.0

        # Similitud por contenido
        if name1_base in name2_base or name2_base in name1_base:
            return 0.8

        # Similitud por caracteres comunes (algoritmo simple)
        common_chars = set(name1_base) & set(name2_base)
        total_chars = set(name1_base) | set(name2_base)

        if not total_chars:
            return 0.0

        return len(common_chars) / len(total_chars)




